"""
专门用于在python-docx读取docx文档之前的一些解析和转换操作，防止后续报错
"""
import io
import os
import re
import subprocess
from docx import Document


def detect_docx_vendor_and_version(xml_bytes: bytes):
    """
    检测DOCX文档厂商和版本
    Return: ("厂商", 主要版本号(int))
    """
    xml_str = xml_bytes.decode('utf-8', errors='ignore')
    print(xml_str)

    # re匹配Application
    app_match = re.search(r'<Application>([^<]+)</Application>', xml_str)
    app_name = app_match.group(1) if app_match else None
    print(app_match)
    print(app_name)

    # AppVersion
    version_match = re.search(r'<AppVersion>([^<]+)</AppVersion>', xml_str)
    version = version_match.group(1) if version_match else None
    print(version)

    vendor = '未知厂商，没有Application标记'
    if app_name:
        if 'Microsoft' in app_name:
            vendor = 'Microsoft'
        elif 'WPS' in app_name or 'Kingsoft' in app_name:
            vendor = 'WPS'
        elif 'LibreOffice' in app_name:
            vendor = 'LibreOffice'
        elif 'OpenOffice' in app_name:
            vendor = 'OpenOffice'
        elif 'Apple' in app_name or 'Pages' in app_name:
            vendor = 'Apple'
        elif 'Apache' in app_name:  # 手机生成的
            vendor = 'Apache'

    major_version = 0
    if version:
        try:
            major_version = int(version.split('.')[0])
            print("版本", major_version)
        except (ValueError, AttributeError):
            major_version = 0

    return vendor, major_version


def convert_doc_to_docx(file_path, output_file):
    """用 LibreOffice 将 .doc 转换为 .docx"""
    try:
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "docx",
            "--outdir",
            os.path.dirname(output_file),
            file_path
        ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return output_file if os.path.exists(output_file) else None
    except Exception as e:
        print(f"LibreOffice 转换失败: {e}")
        return None


def process_doc(binary_stream: bytes, file_extension=None, stop_flag=None, progress=None, bizType=None):
    temp_doc_path = './temp/temp.doc'
    temp_docx_path = './temp/temp.docx'

    # 确保临时目录存在
    os.makedirs('./temp', exist_ok=True)

    header = binary_stream[:20]
    doc = None

    try:
        # Word 97-2003 格式
        if header.startswith(b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1'):
            print("检测到为word97-2003，修复中……")
            with open(temp_doc_path, 'wb') as f:
                f.write(binary_stream)
            if not convert_doc_to_docx(temp_doc_path, temp_docx_path):
                raise Exception("Word 97-2003文件转换失败")

            with open(temp_docx_path, 'rb') as f:
                stream = io.BytesIO(f.read())
            doc = Document(stream)

        # OpenXML 格式 (DOCX)
        elif header.startswith(b'PK\x03\x04'):
            import zipfile
            stream = io.BytesIO(binary_stream)

            with zipfile.ZipFile(stream, 'r') as docx:
                if 'docProps/app.xml' not in docx.namelist():
                    print("警告: 未找到docProps/app.xml，直接解析")
                    stream = io.BytesIO(binary_stream)
                    doc = Document(stream)
                else:
                    app_xml = docx.read('docProps/app.xml')
                    vendor, major_version = detect_docx_vendor_and_version(app_xml)

                    print(f"检测到厂商: {vendor}, 主要版本: {major_version}")

                    # 只对Microsoft文档进行版本判断
                    if major_version < 15:  # Word 2013以下
                        print(f"文档版本 {major_version} 低于15，需要转换")
                        with open(temp_doc_path, 'wb') as f:
                            f.write(binary_stream)
                        if not convert_doc_to_docx(temp_doc_path, temp_docx_path):
                            raise Exception("低版本文档转换失败")
                        with open(temp_docx_path, 'rb') as f:
                            stream = io.BytesIO(f.read())
                        doc = Document(stream)
                    else:
                        # 版本足够高，直接解析
                        stream = io.BytesIO(binary_stream)
                        doc = Document(stream)
        else:
            raise Exception(f"不支持的文件格式: {file_extension}")

        return doc, "", ""

    except Exception as e:
        print(f'error {e}')
        return None, f"文件解析失败: {str(e)}", ""
    finally:
        print('delete temp file')
        # 清理临时文件
        for path in [temp_doc_path, temp_docx_path]:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception as cleanup_error:
                print(f"清理临时文件失败 {path}: {cleanup_error}")


if __name__ == '__main__':
    # test_docx = r"./docs/test/中间版本-赢他客服平台系统采购-招标文件V3.docx"
    # test_docx = r"./docs/errors/小图标-富润标书-天翼云-合稿 v3 1030.docx"
    test_docx = r'./docs/test/黄岩区中心粮库“浙江粮仓”智慧粮库升级改造项目资信技术文件-1013改错-qwen3-235B.docx'
    with open(test_docx, 'rb') as f:
        binary_stream = f.read()
        doc, error_msg, table_msg = process_doc(binary_stream)
        if doc:
            print("文档解析成功")
        else:
            print(f"文档解析失败: {error_msg}")